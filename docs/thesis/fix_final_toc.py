from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "MindHelper_diploma_final.docx"


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_outline_level(paragraph: Paragraph, level: int | None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:outlineLvl"))
    if existing is not None:
        p_pr.remove(existing)
    if level is not None:
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level))
        p_pr.append(outline)


def add_field_run(paragraph: Paragraph, field_type: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), field_type)
    run._r.append(fld_char)


def add_toc_field(paragraph: Paragraph) -> None:
    add_field_run(paragraph, "begin")

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    instr_run._r.append(instr)

    add_field_run(paragraph, "separate")
    paragraph.add_run("Оглавление обновляется средствами Microsoft Word.")
    add_field_run(paragraph, "end")


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def is_level_0(text: str) -> bool:
    if text in {
        "Определения, обозначения и сокращения",
        "Введение",
        "Заключение",
        "Список использованных источников",
    }:
        return True
    if re.match(r"^[1-4]\s", text):
        return True
    if text.startswith("Приложение "):
        return True
    return False


def is_level_1(text: str) -> bool:
    return bool(re.match(r"^[1-4]\.\d+\s", text))


def main() -> None:
    document = Document(TARGET)

    toc_entry = next(((i, p) for i, p in enumerate(document.paragraphs) if p.text.strip() == "Содержание"), None)
    if toc_entry is None:
        raise RuntimeError("Heading 'Содержание' not found")
    idx, toc_heading = toc_entry

    paragraphs = list(document.paragraphs)
    if idx + 1 < len(paragraphs):
        delete_paragraph(paragraphs[idx + 1])
    toc_paragraph = insert_after(toc_heading)
    add_toc_field(toc_paragraph)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if is_level_0(text):
            set_outline_level(paragraph, 0)
        elif is_level_1(text):
            set_outline_level(paragraph, 1)
        else:
            set_outline_level(paragraph, None)

    set_update_fields(document)
    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
