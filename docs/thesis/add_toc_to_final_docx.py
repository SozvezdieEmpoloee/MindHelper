from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "MindHelper_diploma_final.docx"


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_field_run(paragraph: Paragraph, field_type: str, text: str | None = None) -> None:
    run = paragraph.add_run()
    r = run._r
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), field_type)
    r.append(fld_char)
    if text:
        text_el = OxmlElement("w:t")
        text_el.text = text
        r.append(text_el)


def add_toc_field(paragraph: Paragraph) -> None:
    add_field_run(paragraph, "begin")

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\t "Введение/Заключение,1,Глава,1,Параграф,2" \\h \\z \\u'
    instr_run._r.append(instr)

    add_field_run(paragraph, "separate")
    paragraph.add_run("Оглавление обновляется средствами Microsoft Word.")
    add_field_run(paragraph, "end")


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main() -> None:
    document = Document(TARGET)
    toc_entry = next(((i, p) for i, p in enumerate(document.paragraphs) if p.text.strip() == "Содержание"), None)
    if toc_entry is None:
        raise RuntimeError("Heading 'Содержание' not found")
    idx, toc_heading = toc_entry

    next_paragraph = None
    paragraphs = list(document.paragraphs)
    if idx + 1 < len(paragraphs):
        next_paragraph = paragraphs[idx + 1]
    if next_paragraph and next_paragraph.text.strip().startswith("Содержание формируется"):
        delete_paragraph(next_paragraph)

    toc_paragraph = insert_after(toc_heading)
    add_toc_field(toc_paragraph)
    set_update_fields(document)
    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
