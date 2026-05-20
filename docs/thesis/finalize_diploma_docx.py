from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MindHelper_diploma_v4_corrected.docx"
TARGET = ROOT / "MindHelper_diploma_v5_final.docx"

STYLE_BODY = "Основной текст курсовой"
STYLE_INTRO = "Введение/Заключение"
STYLE_CHAPTER = "Глава"
STYLE_SECTION = "Параграф"
STYLE_CAPTION = "Название рисунка"
STYLE_LIST = "Списки"
STYLE_CODE = "Код"
STYLE_SOURCE = "Список источников"


def insert_after(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.style = style
    p.add_run(text)
    return p


def insert_block_after(document: Document, heading: str, block: list[str]) -> None:
    anchor = next((p for p in document.paragraphs if p.text.strip() == heading), None)
    if anchor is None:
        return
    current = anchor
    for text in block:
        current = insert_after(current, text, STYLE_BODY)


def replace_text_everywhere(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text
        for old, new in replacements.items():
            text = text.replace(old, new)
        if text != paragraph.text:
            paragraph.clear()
            paragraph.add_run(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for old, new in replacements.items():
                        text = text.replace(old, new)
                    if text != paragraph.text:
                        paragraph.clear()
                        paragraph.add_run(text)


def apply_numbering_fixes(document: Document) -> None:
    replace_text_everywhere(
        document,
        {
            "Упрощенное устройство выбранной модели показано на рисунке 2.3.": "Упрощенное устройство выбранной модели показано на рисунке 1.1.",
            "Рисунок 2.3 - Устройство Qwen под капотом": "Рисунок 1.1 - Устройство Qwen3 под капотом",
            "2.5 Конвейер формирования ответа": "2.4 Конвейер формирования ответа",
            "2.6 Safety-flow и профиль состояния": "2.5 Safety-flow и профиль состояния",
            "2.7 Аудит safety-решений": "2.6 Аудит safety-решений",
            "2.8 Формальная модель маршрутизации": "2.7 Формальная модель маршрутизации",
            "2.9 Red-team корпус и предопределённые сценарии": "2.8 Red-team корпус и предопределённые сценарии",
            "2.10 Методика safety evaluation": "2.9 Методика safety evaluation",
            "рисунке 2.4": "рисунке 2.3",
            "Рисунок 2.4 - Конвейер формирования ответа нейросети": "Рисунок 2.3 - Конвейер формирования ответа нейросети",
            "рисунке 2.5": "рисунке 2.4",
            "Рисунок 2.5 - Safety-flow и профиль состояния": "Рисунок 2.4 - Safety-flow и профиль состояния",
            "таблице 2.3": "таблице 2.2",
            "Таблица 2.3 - Уровни риска и действия системы": "Таблица 2.2 - Уровни риска и действия системы",
            "таблице 2.4": "таблице 2.3",
            "Таблица 2.4 - Поля решения safety-маршрутизатора": "Таблица 2.3 - Поля решения safety-маршрутизатора",
            "таблице 2.5": "таблице 2.4",
            "Таблица 2.5 - Различие red-team и предопределённые сценарии": "Таблица 2.4 - Различие red-team корпуса и предопределённых сценариев",
            "Таблица 2.5 - Различие red-team корпуса и предопределённых сценариев": "Таблица 2.4 - Различие red-team корпуса и предопределённых сценариев",
            "таблице 2.6": "таблице 2.5",
            "Таблица 2.6 - Критерии оценки безопасности": "Таблица 2.5 - Критерии оценки безопасности",
            "Предопределнные": "Предопределённые",
            "предопределнные": "предопределённые",
        },
    )


def add_extra_text(document: Document) -> None:
    insert_block_after(document, "1.2 Существующие цифровые решения", [
        "Критический анализ источников показывает, что открытым остается вопрос соединения трех требований: естественного диалога, локального контроля данных и формализованной безопасности. Многие публикации и продуктовые решения рассматривают чат-бота как самостоятельного помощника, но недостаточно подробно описывают внутреннюю маршрутизацию риска. В данной работе этот пробел закрывается через построение наблюдаемой архитектуры, где каждый ответ связан с маршрутом обработки.",
        "Таким образом, место разрабатываемой системы в предметной области определяется не конкуренцией с профессиональными платформами телемедицины, а созданием инженерно прозрачного прототипа предварительной поддержки. Такой прототип может использоваться для исследования поведения моделей, проверки safety-сценариев и подготовки последующего датасета для улучшения нейросетевого компонента.",
    ])
    insert_block_after(document, "1.3 Риски применения LLM в психологическом диалоге", [
        "Особую сложность представляет риск ложного спокойствия. Модель может сформировать мягкий и внешне заботливый ответ, но при этом не распознать непосредственную угрозу. Для пользователя такой ответ опасен тем, что создает ощущение нормальности ситуации и не направляет к срочной помощи. Поэтому в проекте приоритетом является снижение false negative для critical-сценариев, даже если это увеличивает число осторожных срабатываний.",
        "Другой важный риск связан с галлюцинацией справочной информации. В психологическом сервисе недопустимо, чтобы модель выдумывала телефон горячей линии, адрес клиники или имя специалиста. По этой причине справочные данные отделены от LLM и хранятся в таблицах emergency_resource, specialist и specialist_location. Модель может участвовать в формулировке поддерживающего ответа, но фактические ресурсы помощи должны извлекаться из базы данных.",
    ])
    insert_block_after(document, "2.2 Последовательность обработки сообщения", [
        "Оценка сообщения не сводится к элементарному списку ключевых слов. В проекте учитывается композиция признаков: личное местоимение, упоминание вреда, наличие способа, срочность, подготовленность, просьба о вредной инструкции и контекст третьего лица. Чем больше признаков указывает на непосредственность и личную направленность риска, тем выше уровень маршрутизации. Это позволяет срабатывать не только на заранее известные фразы, но и на новые формулировки.",
        "При отсутствии критических признаков система не должна превращаться в жесткий фильтр. В low и elevated-сценариях важна полезность: пользователь ожидает не только сочувствия, но и конкретного шага. Поэтому после оценки риска выбирается режим ответа: практический совет, поддерживающий план, уточнение состояния или отражение переживания. Такой режим передается в neural_engine и влияет на системную инструкцию для модели.",
    ])
    insert_block_after(document, "3.8 Реализация API и сервисного слоя", [
        "CRUD-операции в проекте распределены по ролям. Пользователь создает сообщения и читает собственную историю, администратор управляет справочниками, версиями модели и модерационными сущностями, а системные сервисы создают crisis_event и safety_audit_log автоматически. Такое разделение снижает риск случайного изменения критичных данных через пользовательский интерфейс.",
        "Сервисный слой также упрощает дальнейшее расширение. Если будет добавлен мобильный клиент или другой мессенджер, он сможет использовать тот же сценарий обработки сообщений. Для дипломной работы это важно, поскольку показывает, что реализован не набор отдельных экранов, а повторно используемая архитектура вопросно-ответной системы.",
    ])
    insert_block_after(document, "4.1 Функциональное тестирование", [
        "Фрагменты кода, иллюстрирующие хранение safety-аудита, обработку одного сообщения и работу маршрутизатора, приведены в приложениях Б-Г. Они показывают, что тестируемая логика находится не только в представлениях API, но и в отдельных сервисах, пригодных для модульной проверки.",
    ])
    insert_block_after(document, "4.3 Оценка модели Qwen3", [
        "Оценка Qwen3 проводилась после выбора архитектурной рамки безопасности. Это важно, потому что модель не должна оцениваться как самостоятельный психологический консультант. В эксперименте проверялось, насколько хорошо она работает внутри заданных ограничений: отвечает на русском языке, следует системной инструкции, дает практические шаги в безопасных сценариях и не пытается обойти запреты.",
        "Для обычных запросов оценивалась способность модели не уходить в пустую эмпатию. Например, на запрос о расслаблении ожидается не только фраза о понимании состояния, но и короткая последовательность действий: дыхание, снижение стимулов, вода, небольшой бытовой шаг, план на ближайшие минуты. Для elevated-сценариев дополнительно оценивалась осторожность формулировок и рекомендация обратиться к специалисту при ухудшении состояния.",
        "Для high и critical-сценариев качество свободного ответа модели не является главным критерием, поскольку в таких ветках генерация ограничивается или запрещается. Главным становится способность всей системы не передать опасный запрос в обычный генеративный режим. Поэтому оценка Qwen3 связана с оценкой prompt-policy, response policy и маршрутизатора.",
    ])
    insert_block_after(document, "4.4 Ограничения", [
        "Ограничения системы связаны с тем, что прототип не проходил клиническую валидацию и не должен использоваться как медицинский инструмент. Даже при корректной маршрутизации он не заменяет очную консультацию, диагностику или экстренную помощь. Это ограничение должно быть явно отражено в пользовательском интерфейсе и в сообщениях Telegram-бота.",
        "Техническое ограничение состоит в зависимости качества ответа от локальной модели и аппаратных ресурсов. Более крупные модели могут давать более сильный диалог, но требуют значительно больше памяти и вычислительной мощности. Поэтому в работе выбран компромиссный вариант: локально запускаемая модель, дополненная строгим safety-контуром и возможностью последующего обновления версии.",
        "Методическое ограничение связано с red-team корпусом. Невозможно заранее перечислить все формулировки, которыми пользователь может описать кризис. Поэтому корпус должен развиваться итеративно: после тестирования, анализа ошибок и экспертной разметки новых примеров. Это направление напрямую связано с перспективой формирования датасета для будущего дообучения.",
    ])


def fix_styles(document: Document) -> None:
    section_titles = {
        "ER-диаграмма базы данных",
        "Фрагмент кода модели safety-аудита",
        "Фрагмент кода обработки одного сообщения",
        "Фрагмент кода safety-маршрутизатора",
    }
    in_sources = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "Список использованных источников":
            in_sources = True
            paragraph.style = STYLE_INTRO
        elif text.startswith("Приложение "):
            paragraph.style = STYLE_CHAPTER
        elif text in section_titles:
            paragraph.style = STYLE_SECTION
        elif paragraph.style.name in {STYLE_INTRO, STYLE_CHAPTER, STYLE_SECTION, STYLE_CAPTION, STYLE_LIST, STYLE_CODE, STYLE_SOURCE, STYLE_BODY}:
            pass
        elif text.startswith("Рисунок ") or text.startswith("Таблица "):
            paragraph.style = STYLE_CAPTION
        elif in_sources:
            paragraph.style = STYLE_SOURCE
        else:
            paragraph.style = STYLE_BODY

        for run in paragraph.runs:
            if paragraph.style.name == STYLE_CODE:
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
                run.font.size = Pt(9)
            else:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(14)

        if paragraph.style.name == STYLE_CAPTION:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)


def fix_sections_and_images(document: Document) -> None:
    for section in document.sections:
        if section.page_width and section.page_height and section.orientation == WD_ORIENTATION.PORTRAIT:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0 if section.orientation == WD_ORIENTATION.PORTRAIT else 2.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    for idx, shape in enumerate(document.inline_shapes):
        max_width = Cm(16.0)
        if idx == len(document.inline_shapes) - 1:
            max_width = Cm(25.0)
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)


def main() -> None:
    shutil.copyfile(SOURCE, TARGET)
    document = Document(TARGET)
    apply_numbering_fixes(document)
    add_extra_text(document)
    fix_styles(document)
    fix_sections_and_images(document)
    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
