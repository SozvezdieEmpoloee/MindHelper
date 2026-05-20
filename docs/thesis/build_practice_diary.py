from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Dnevnik_NIR_KretovDV_2026.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def format_paragraph(paragraph, *, align=None, bold=False, size=14, first_line=True) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_line else Cm(0)
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.font.bold = bold


def add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    run.font.bold = True


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    format_paragraph(p)


def add_signature_line(document: Document, left: str, right: str = "") -> None:
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(left)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    if right:
        run = p.add_run("\t" + right)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [0.9, 2.6, 5.6, 7.4]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths[col_idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
                    run.font.bold = row_idx == 0
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")


def add_calendar_table(document: Document) -> None:
    add_body(
        document,
        "Календарный план составлен по этапам выполнения выпускной квалификационной работы на тему "
        "«Разработка вопросно-ответной системы для предварительной психологической диагностики на основе нейросети». "
        "Период практики охватывает февраль, март и апрель 2026 года.",
    )
    caption = document.add_paragraph()
    caption.add_run("Таблица 1 - Календарный план работы за период 01.02.2026 - 30.04.2026")
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(caption, first_line=False)

    rows = [
        (
            "1",
            "01.02 - 09.02",
            "Уточнение темы ВКР, анализ предметной области, формирование требований к сервису психологической поддержки.",
            "Сформулированы цель, объект, предмет и задачи работы. Определены ключевые ограничения: сервис не ставит диагноз, не заменяет специалиста и должен иметь отдельный контур безопасности.",
        ),
        (
            "2",
            "10.02 - 18.02",
            "Проектирование ER-модели и структуры базы данных.",
            "Разработана компактная ER-диаграмма с UUID-идентификаторами. Выделены сущности пользователей, чатов, сообщений, кризисных событий, справочника специалистов, экстренных ресурсов, версий нейросети и административного аудита.",
        ),
        (
            "3",
            "19.02 - 29.02",
            "Создание PostgreSQL-базы данных, подготовка SQL-скриптов и тестовых данных.",
            "Подготовлены скрипты создания сущностей и связей, добавлены тестовые записи пользователей, специалистов, организаций помощи, экстренных телефонов и версии нейросетевой модели.",
        ),
        (
            "4",
            "01.03 - 10.03",
            "Разработка backend-части на Django и Django REST Framework.",
            "Создана структура приложений backend, реализованы модели, сериализаторы, API, CRUD-операции, регистрация, авторизация, хранение истории диалога и базовая административная панель.",
        ),
        (
            "5",
            "11.03 - 19.03",
            "Интеграция локальной языковой модели Qwen3 через Ollama.",
            "Проведено сравнение вариантов локального запуска, выбрана Qwen3 как практичный вариант для русскоязычного диалога. Реализован сервис обращения к Ollama и сохранение сведений о версии модели.",
        ),
        (
            "6",
            "20.03 - 31.03",
            "Разработка Telegram-бота как дополнительного канала взаимодействия.",
            "Реализованы команды запуска, справки, приватности и экстренных ресурсов. Telegram-бот подключен к общему backend-контуру, чтобы сообщения проходили через единый сервис чата и общую базу данных.",
        ),
        (
            "7",
            "01.04 - 10.04",
            "Построение safety-flow и сценариев обработки психологически чувствительных сообщений.",
            "Реализованы уровни риска low, elevated, high и critical, ASQ-подобный уточняющий сценарий, policy layer, запрет свободной генерации в critical-сценариях и показ экстренных ресурсов из базы данных.",
        ),
        (
            "8",
            "11.04 - 20.04",
            "Доработка frontend-части и административного интерфейса.",
            "Улучшены пользовательские блоки сайта, чат, каталог специалистов, карта, описание сервиса. В Django Admin добавлены возможности управления версиями модели, экстренными ресурсами, справочниками и журналом safety-аудита.",
        ),
        (
            "9",
            "21.04 - 30.04",
            "Тестирование, анализ результатов и подготовка материалов ВКР.",
            "Проведены функциональные и safety-тесты: регистрация, уникальность email, chat API, работа Telegram-бота, маршрутизация риска и аудит. Подготовлены диаграммы, таблицы, приложения и текст практической части дипломной работы.",
        ),
    ]
    table = document.add_table(rows=1, cols=4)
    headers = ["№", "Период", "Планируемые мероприятия", "Фактическое выполнение"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
    style_table(table)


def add_research_work(document: Document) -> None:
    add_heading(document, "2 Научно-исследовательская работа")
    paragraphs = [
        "Научно-исследовательская работа была направлена на разработку и исследование вопросно-ответной системы предварительной психологической диагностики и поддержки пользователя на основе нейросети. В качестве центральной задачи рассматривалось не только создание чат-интерфейса, но и построение безопасной архитектуры, в которой языковая модель не принимает критические решения без дополнительного программного контроля.",
        "На первом этапе был проведен анализ предметной области. Были изучены особенности цифровых сервисов психологической поддержки, ограничения применения больших языковых моделей в чувствительном диалоге, риски ложного успокоения, вредных рекомендаций и галлюцинации справочной информации. На основе анализа была сформулирована роль системы: она предназначена для предварительной поддержки, самонаблюдения и маршрутизации пользователя к ресурсам помощи, но не является медицинским диагностическим инструментом.",
        "Далее была спроектирована модель данных. В ER-модель включены сущности для хранения учетных записей пользователей, каналов взаимодействия, единого пользовательского чата, сообщений, кризисных событий, справочника специалистов, адресов, экстренных ресурсов, версий нейросети и журнала safety-аудита. Использование UUID позволяет избежать зависимости от последовательных числовых идентификаторов и упрощает дальнейшее масштабирование сервиса.",
        "Практическая часть включала реализацию backend на Django и Django REST Framework. Были разработаны модели, миграции, сериализаторы, представления API, сервисный слой, регистрация и авторизация пользователей, CRUD-операции, административные формы и тесты. Важной особенностью реализации стало отделение бизнес-логики от представлений: обработка одного сообщения, маршрутизация риска и запись аудита вынесены в сервисы, которые могут использоваться как сайтом, так и Telegram-ботом.",
        "В рамках исследования нейросетевого компонента была выбрана локальная модель Qwen3, запускаемая через Ollama. Такой подход позволяет не передавать пользовательские сообщения во внешний коммерческий API и сохраняет контроль над версией модели. В базе данных хранится объект версии нейросети, что дает возможность документировать применяемую модель, ее provider, профиль безопасности и параметры генерации.",
        "Отдельным результатом стала разработка safety-flow. Система разделяет сообщения на уровни low, elevated, high и critical. Для low и elevated допускается нейросетевая генерация с ограничениями prompt-policy и response policy. Для high-сценариев используется уточняющий сценарий оценки риска. Для critical-сценариев свободная генерация блокируется, а пользователь получает предопределенный ответ с экстренными ресурсами, загружаемыми из базы данных.",
        "Также был реализован Telegram-бот как дополнительный канал доступа. Он не содержит отдельной независимой логики консультации, а использует общий backend и те же сущности базы данных. Это позволяет сохранять единый принцип работы системы: независимо от канала взаимодействия пользовательские сообщения проходят через общий сервис чата, safety-flow и журнал аудита.",
        "В ходе тестирования проверялись регистрация и запрет повторного email, работа API сообщений, сохранение истории чата, создание кризисных событий, получение экстренных ресурсов, маршрутизация dangerous-сценариев, управление моделью через административную панель и корректность интеграции Telegram-бота. Дополнительно был сформирован red-team корпус сценариев для проверки устойчивости системы к прямым и косвенным опасным формулировкам.",
        "По итогам практики подготовлены материалы для выпускной квалификационной работы: описание предметной области, ER-диаграмма, архитектурные диаграммы, таблицы требований, описание реализации, фрагменты кода приложений, описание тестирования, оценка выбранной модели и направления дальнейшего развития. Полученный прототип MindHelper является основой для дальнейшего улучшения нейросетевого поведения, расширения корпуса сценариев и подготовки обезличенного датасета с экспертной разметкой.",
    ]
    for paragraph in paragraphs:
        add_body(document, paragraph)

    document.add_paragraph()
    add_signature_line(document, "Обучающийся ____________________ Кретов Д.В.")
    add_signature_line(document, "Руководитель ____________________ Шишко Ю.В.")


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    styles = document.styles
    for style_name in ["Normal", "Table Grid"]:
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading(document, "1 Календарный план работы с февраля по апрель")
    add_calendar_table(document)
    document.add_paragraph()
    add_research_work(document)
    document.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
