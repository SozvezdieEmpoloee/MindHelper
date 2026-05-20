from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "08_diploma_v2_extended.md"
TARGET_MD = ROOT / "09_diploma_v3.md"
TARGET_DOCX = ROOT / "MindHelper_diploma_v3.docx"

FIG_PREFIX = "\u0420\u0438\u0441\u0443\u043d\u043e\u043a "
TABLE_PREFIX = "\u0422\u0430\u0431\u043b\u0438\u0446\u0430 "
PLACEHOLDER_WORD = "\u041c\u0435\u0441\u0442\u043e"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B8C1CC")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x16, 0x2A, 0x44)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.2
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.size = Pt(14)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(3)


def add_inline_text(paragraph, text: str, size: int = 14) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        is_code = len(part) >= 2 and part.startswith("`") and part.endswith("`")
        run = paragraph.add_run(part[1:-1] if is_code else part)
        run.font.name = "Courier New" if is_code else "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
        run.font.size = Pt(max(size - 1, 9) if is_code else size)
        if is_code:
            run.font.color.rgb = RGBColor(0x27, 0x3B, 0x4F)


def add_paragraph(document: Document, text: str, style: str | None = None, size: int = 14):
    paragraph = document.add_paragraph(style=style)
    add_inline_text(paragraph, text, size=size)
    return paragraph


def add_caption(document: Document, text: str):
    paragraph = add_paragraph(document, text, size=12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(9)
    for run in paragraph.runs:
        run.italic = True
    return paragraph


def add_placeholder(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, "F2F6FA")
    set_cell_margins(cell, 220, 220, 220, 220)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text.strip("[]"))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0x4D, 0x5F, 0x73)
    document.add_paragraph()


def split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def is_alignment_row(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if line.strip()]
    rows = [row for row in rows if not all(re.fullmatch(r":?-{3,}:?", c) for c in row)]
    if not rows:
        return

    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "DCEAF7")
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_text(paragraph, text, size=10 if cols >= 4 else 11)
            for run in paragraph.runs:
                if r_idx == 0:
                    run.bold = True
    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F5F7")
    p_pr.append(shd)
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)
        if idx != len(lines) - 1:
            run.add_break()


def add_image(document: Document, md_path: str) -> None:
    image_path = (ROOT / md_path).resolve()
    if not image_path.exists():
        add_placeholder(document, f"Image file is not available: {md_path}")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=Cm(15.0))
    except Exception:
        run.add_picture(str(image_path), width=Cm(13.5))


def build_docx() -> None:
    shutil.copyfile(SOURCE_MD, TARGET_MD)
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    configure_document(document)

    first_h1_seen = False
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip()

        if in_code:
            if line.startswith("```"):
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.startswith("```"):
            in_code = True
            code_lines = []
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            heading_text = line[level:].strip()
            if level == 1:
                if first_h1_seen:
                    document.add_section(WD_SECTION.NEW_PAGE)
                first_h1_seen = True
                paragraph = document.add_heading(heading_text, level=1)
            else:
                paragraph = document.add_heading(heading_text, level=min(level, 3))
            paragraph.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        if re.match(r"!\[.*?\]\((.*?)\)", line):
            md_path = re.match(r"!\[.*?\]\((.*?)\)", line).group(1)
            add_image(document, md_path)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and is_alignment_row(lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(document, table_lines)
            continue

        if line.startswith(FIG_PREFIX) or line.startswith(TABLE_PREFIX):
            add_caption(document, line)
            i += 1
            continue

        if line.startswith("[") and line.endswith("]") and PLACEHOLDER_WORD in line:
            add_placeholder(document, line)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", line)
        number = re.match(r"^\d+\.\s+(.*)$", line)
        if bullet:
            paragraph = add_paragraph(document, bullet.group(1), style="List Bullet")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue
        if number:
            paragraph = add_paragraph(document, number.group(1), style="List Number")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        paragraph = add_paragraph(document, line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1

    if in_code and code_lines:
        add_code_block(document, code_lines)

    document.save(TARGET_DOCX)


if __name__ == "__main__":
    build_docx()
    print(TARGET_MD)
    print(TARGET_DOCX)
