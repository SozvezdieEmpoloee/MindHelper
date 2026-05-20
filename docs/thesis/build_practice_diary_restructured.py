from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'Dnevnik_NIR_KretovDV_2026_restructured.docx'


def set_font(run, size=14, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_end)


def para(doc, text='', *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True, size=14, bold=False, after=4, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    set_font(r, size, bold)
    return p


def heading(doc, text):
    p = para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, first=False, bold=True, after=6, before=8)
    return p


def setup_doc():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(3)
    s.right_margin = Cm(1.5)
    add_page_number(s)
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(14)
    return doc


def title_page(doc):
    for line in [
        'МИНОБРНАУКИ РОССИИ',
        'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ',
        'ВЫСШЕГО  ОБРАЗОВАНИЯ',
        '«ВОРОНЕЖСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»',
        '(ФГБОУ ВО «ВГУ»)',
        'Факультет компьютерных наук',
        'Кафедра программирования и информационных технологий',
    ]:
        para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, after=0)
    doc.add_paragraph('\n\n')
    para(doc, 'ДНЕВНИК', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True, after=0)
    para(doc, 'по производственной практике', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True, after=0)
    para(doc, '(научно-исследовательская работа)', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True, after=10)
    para(doc, '8 семестр', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, after=22)
    para(doc, '09.03.02 Информационные системы и технологии', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, after=0)
    para(doc, 'Программная инженерия в информационных системах', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, after=20)
    para(doc, 'Тема: «Разработка вопросно-ответной системы для предварительной психологической диагностики на основе нейросети»', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, after=20)
    para(doc, 'Обучающийся ____________________ Кретов Д.В., 4 курс, д/о', align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, after=0)
    para(doc, 'Руководитель ____________________ Шишко Ю.В., ассистент', align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, after=0)
    doc.add_paragraph('\n\n')
    para(doc, 'Воронеж 2026', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, after=0)
    doc.add_paragraph()


def calendar_table(doc):
    heading(doc, '1. Календарный план работы')
    para(doc, 'Таблица 1 – Календарный план работы (Период: 10.02.2026 – 03.05.2026)', align=WD_ALIGN_PARAGRAPH.CENTER, first=False, size=12, after=4)
    rows = [
        ('10.02', 'Ознакомление с задачами производственной практики и уточнение темы выпускной работы', 'Актуализированы цель, задачи и план разработки'),
        ('17.02', 'Анализ предметной области предварительной психологической поддержки и рисков применения LLM', 'Сформированы требования к сервису и safety-контуре'),
        ('24.02', 'Проектирование ER-модели базы данных и основных сущностей сервиса', 'Определены таблицы пользователей, чата, сообщений, модели и событий риска'),
        ('03.03', 'Создание PostgreSQL-базы данных и подготовка тестовых данных', 'Разработана структура БД и заполнены справочники'),
        ('10.03', 'Реализация backend на Django REST Framework', 'Созданы модели, сериализаторы, API и сервисный слой'),
        ('17.03', 'Разработка web-интерфейса пользователя и сценариев работы с чатом', 'Реализованы регистрация, авторизация и базовый чат'),
        ('24.03', 'Интеграция Qwen3 через Ollama и фиксация версии модели в базе данных', 'Подключён локальный нейросетевой компонент'),
        ('31.03', 'Разработка safety-flow и маршрутизации сообщений по уровням риска', 'Реализованы low, elevated, high и critical сценарии'),
        ('07.04', 'Реализация Telegram-бота как дополнительного канала взаимодействия', 'Бот подключён к общему chat-сервису backend'),
        ('14.04', 'Добавление справочников экстренной помощи и организаций для отображения пользователю', 'Заполнены emergency_resource и данные для карты'),
        ('21.04', 'Проведение функционального, интеграционного и safety-тестирования', 'Проверены API, чат, Telegram-бот и кризисные сценарии'),
        ('28.04', 'Подготовка материалов выпускной работы и отчёта по практике', 'Сформированы текстовые разделы, таблицы и диаграммы'),
        ('03.05', 'Итоговый анализ результатов разработки и определение направлений развития', 'Сформулированы выводы и перспективы проекта'),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    headers = ['Дата', 'Содержание выполненной работы', 'Результат']
    for i, h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[i], h, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(tbl.rows[0].cells[i], 'D9EAF7')
    for date, work, result in rows:
        cells = tbl.add_row().cells
        set_cell_text(cells[0], date, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], work, size=10.5)
        set_cell_text(cells[2], result, size=10.5)
    widths = [1.55, 8.1, 5.15]
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)
    doc.add_page_break()


def description(doc):
    heading(doc, '2. Краткое описание выполненной работы')
    heading(doc, '10.02 – 03.03')
    para(doc, 'Работа начата с уточнения задач производственной практики и анализа ранее подготовленных материалов по теме выпускной квалификационной работы. Были определены ключевые ограничения предметной области: необходимость предварительной психологической поддержки, недопустимость медицинской диагностики средствами сервиса и обязательное наличие safety-контура для опасных сценариев.')
    para(doc, 'На данном этапе выполнено проектирование структуры данных. В ER-модель включены сущности user_account, channel_account, user_chat, chat_message, neural_model_version, crisis_event, emergency_resource и связанные справочники. Это позволило заранее разделить хранение истории диалога, версий нейросетевой модели, событий риска и проверенных ресурсов помощи.')

    heading(doc, '03.03 – 24.03')
    para(doc, 'Следующий этап был связан с переходом от проектирования к программной реализации. На базе PostgreSQL создана структура базы данных и подготовлены тестовые записи. В backend на Django REST Framework реализованы модели, миграции, сериализаторы, представления API и сервисный слой обработки сообщений.')
    para(doc, 'Параллельно был доработан web-интерфейс. Для пользователя реализованы регистрация, авторизация и чат поддержки. Основной принцип реализации состоит в том, что frontend не содержит собственной логики консультации, а обращается к backend, где выполняются сохранение истории, оценка риска и вызов нейросетевого компонента.')

    heading(doc, '24.03 – 07.04')
    para(doc, 'На данном этапе выполнена интеграция нейросетевой модели. Для локального запуска выбрана Qwen3, работающая через Ollama. Такой вариант позволяет использовать модель без отправки пользовательских сообщений во внешний API и фиксировать конкретную версию модели в базе данных.')
    para(doc, 'Отдельное внимание уделено safety-flow. Сообщения пользователя маршрутизируются по уровням low, elevated, high и critical. Для обычных сообщений допускается генерация ответа нейросетью, а для критических сценариев свободная генерация должна блокироваться и заменяться безопасным ответом с выводом проверенных ресурсов помощи из базы данных.')

    heading(doc, '07.04 – 21.04')
    para(doc, 'Разработан Telegram-бот как дополнительный канал доступа к сервису. Бот не реализует независимую консультационную логику, а использует общий backend и те же сущности базы данных, что и web-интерфейс. Такой подход обеспечивает единообразное хранение истории сообщений и применение одного safety-flow для разных каналов взаимодействия.')
    para(doc, 'Также были добавлены справочники экстренной помощи и организаций. Для города Воронеж подготовлены данные специалистов и клиник, которые могут отображаться пользователю на сайте. Экстренные контакты хранятся в базе данных, чтобы нейросеть не выдумывала номера телефонов и адреса.')

    heading(doc, '21.04 – 03.05')
    para(doc, 'Завершающий этап включал тестирование и подготовку материалов для выпускной квалификационной работы. Проверялись регистрация, запрет повторного email, работа API сообщений, сохранение истории чата, создание кризисных событий, получение экстренных ресурсов, Telegram-бот и сценарии работы с нейросетью.')
    para(doc, 'Для проверки безопасности использовался red-team набор сообщений, содержащих опасные и пограничные формулировки. Его назначение состоит не в ручной проверке каждого пользователя, а в автоматизированном тестировании маршрутизации риска и предотвращении опасных ответов модели. По итогам работы были сформулированы выводы и направления дальнейшей доработки сервиса.')


def results(doc):
    heading(doc, '3. Итоги выполнения практики')
    para(doc, 'В ходе выполнения производственной практики была разработана программная основа сервиса MindHelper — вопросно-ответной системы для предварительной психологической диагностики и поддержки пользователя на основе нейросети.')
    para(doc, 'Разработанное решение обеспечивает:')
    items = [
        'хранение пользователей, каналов связи и истории диалога в PostgreSQL;',
        'работу web-интерфейса и Telegram-бота через единый backend;',
        'локальную интеграцию Qwen3 через Ollama;',
        'маршрутизацию сообщений по уровням риска в рамках safety-flow;',
        'вывод проверенных экстренных ресурсов из базы данных;',
        'функциональное и safety-тестирование основных сценариев.'
    ]
    for it in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run('− ' + it)
        set_font(r, 14)
    para(doc, 'Полученные результаты подтверждают практическую применимость выбранной архитектуры и создают основу для дальнейшего развития выпускной квалификационной работы. Перспективы включают расширение safety-тестов, улучшение классификации риска и развитие карты специалистов.')
    para(doc, 'Обучающийся ____________________ Кретов Д.В.', align=WD_ALIGN_PARAGRAPH.LEFT, first=False, after=0)
    para(doc, 'Руководитель ____________________ Шишко Ю.В.', align=WD_ALIGN_PARAGRAPH.LEFT, first=False, after=0)


def build():
    doc = setup_doc()
    title_page(doc)
    calendar_table(doc)
    description(doc)
    results(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
